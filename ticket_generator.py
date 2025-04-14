import random
import subprocess
import os
from datetime import datetime
import openai
import requests
from docx import Document as DocxDocument
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
from tkinter import Tk, filedialog, simpledialog, messagebox, ttk, StringVar, Label, Frame, Entry
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


openai.api_key = "KEY"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "2")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tblBorders.append(border)
    tblPr.append(tblBorders)

def read_questions_from_docx(path):
    doc = DocxDocument(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def generate_questions_by_topic(topic, num_questions):
    prompt = (f"Сгенерируй {num_questions} вопросов для экзамена по теме \"{topic}\". "
              f"Верни вопросы в виде списка, по одному вопросу на строке, без нумерации и дополнительного текста.")
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=500
    )
    text = response["choices"][0]["message"]["content"]
    questions = [line.strip() for line in text.split("\n") if line.strip()]
    return questions

def generate_tickets_genetic(questions, num_tickets, generations=50, population_size=100, mutation_rate=0.1):
    if num_tickets * 2 >= len(questions):
        raise ValueError("Недостаточно вопросов для генерации без повторов.")

    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform(questions)
    question_indices = list(range(len(questions)))

    def init_individual():
        indices = question_indices.copy()
        random.shuffle(indices)
        return [indices[i * 2:(i + 1) * 2] for i in range(num_tickets)]

    population = [init_individual() for _ in range(population_size)]

    def ticket_fitness(ticket):
        q1, q2 = ticket
        sim = cosine_similarity(vectors[q1], vectors[q2])[0, 0]
        return 1 - sim

    def individual_fitness(individual):
        return sum(ticket_fitness(ticket) for ticket in individual) / len(individual)

    def crossover(parent1, parent2):
        cp = random.randint(1, num_tickets - 1)
        child = parent1[:cp] + parent2[cp:]
        flat = [q for ticket in child for q in ticket]
        if len(set(flat)) < len(flat):
            missing = list(set(question_indices) - set(flat))
            flat = list(set(flat)) + missing
            random.shuffle(flat)
            child = [flat[i * 2:(i + 1) * 2] for i in range(num_tickets)]
        return child

    def mutate(individual):
        flat = [q for ticket in individual for q in ticket]
        idx1, idx2 = random.sample(range(len(flat)), 2)
        flat[idx1], flat[idx2] = flat[idx2], flat[idx1]
        return [flat[i * 2:(i + 1) * 2] for i in range(num_tickets)]

    for _ in range(generations):
        scored_population = [(ind, individual_fitness(ind)) for ind in population]
        scored_population.sort(key=lambda x: x[1], reverse=True)
        survivors = [ind for ind, fit in scored_population[:max(1, population_size // 5)]]
        next_population = survivors.copy()
        while len(next_population) < population_size:
            parent1, parent2 = random.sample(survivors, 2)
            child = crossover(parent1, parent2)
            if random.random() < mutation_rate:
                child = mutate(child)
            next_population.append(child)
        population = next_population

    best_individual = max(population, key=lambda ind: individual_fitness(ind))
    final_tickets = [[questions[idx] for idx in ticket] for ticket in best_individual]
    return final_tickets

def generate_latex(tickets, output_file, discipline, specialty, group):
    latex_template = r"""
\documentclass[a4paper,12pt]{article}
\usepackage[utf8]{inputenc}
\usepackage[russian]{babel}
\usepackage{geometry}
\geometry{top=2cm, bottom=2cm, left=2cm, right=2cm}
\usepackage{setspace}
\onehalfspacing
\usepackage{array}
\usepackage{times}
\usepackage{multirow}
\usepackage{enumitem}
\setlist[enumerate]{leftmargin=0pt, label=\arabic*.}
\setlength{\arrayrulewidth}{0.5pt}
\begin{document}
{% for ticket in tickets %}
\begin{center}
\begin{tabular}{|p{0.3\textwidth}|p{0.3\textwidth}|p{0.3\textwidth}|} \hline
\parbox{0.3\textwidth}{\centering 
Рассмотрено ПЦК\\
Протокол № \makebox[2cm]{\hrulefill} от \makebox[2cm]{\hrulefill}\\
\makebox[4cm]{\hrulefill} Ф.И.О.\\
председатель ПЦК} &
\parbox{0.3\textwidth}{\centering 
Экзаменационный билет № {{ loop.index }}\\
Дисциплина: {{ discipline }}\\
Специальность: {{ specialty }}\\
Группа: {{ group }}} &
\parbox{0.3\textwidth}{\centering 
УТВЕРЖДАЮ\\
Заместитель директора по УВР \makebox[3cm]{\hrulefill} Ф.И.О.\\
\guillemotleft\makebox[3cm]{\hrulefill} 20\makebox[1cm]{\hrulefill}} \\ \hline
\end{tabular}
\end{center}

\vspace{1cm}

\begin{enumerate}
  \item {{ ticket[0] }}
  \item {{ ticket[1] }}
  \item \hrulefill
\end{enumerate}

\vspace{1cm}

\noindent Преподаватель \hrulefill И.О. Фамилия (подпись)

\newpage
{% endfor %}
\end{document}
    """
    template = Template(latex_template)
    with open(output_file, "w", encoding="utf-8") as f:
        f.write(template.render(tickets=tickets, discipline=discipline, specialty=specialty, group=group))

def compile_pdf(tex_file, output_dir):
    try:
        pdflatex_path = r"C:\Users\artem\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe"
        subprocess.run([pdflatex_path, "-output-directory", output_dir, tex_file], check=True)
        return True
    except subprocess.CalledProcessError:
        return False

def create_formatted_exam_docx(tickets, output_path,
                               discipline="Операционные системы",
                               specialty="02.03.02 ФИиИТ",
                               group="БФИ2102",
                               teacher="И.О. Фамилия",
                               deputy="П.П. Заместитель"):
    doc = DocxDocument()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    for i, ticket in enumerate(tickets, 1):
        header_table = doc.add_table(rows=1, cols=3)
        set_table_borders(header_table)
        widths = [Cm(6), Cm(6), Cm(6)]
        for idx, cell in enumerate(header_table.rows[0].cells):
            cell.width = widths[idx]
        header_table.cell(0, 0).text = (
            "Рассмотрено ПЦК\n"
            "Протокол № ___ от ______\n"
            "__________________ Ф.И.О.\n"
            "председатель ПЦК"
        )
        header_table.cell(0, 1).text = (
            f"Экзаменационный билет № {i}\n"
            f"Дисциплина: {discipline}\n"
            f"Специальность: {specialty}\n"
            f"Группа: {group}"
        )
        header_table.cell(0, 2).text = (
            "УТВЕРЖДАЮ\n"
            "Заместитель директора по УВР\n"
            "____________ " + deputy + "\n"
            "«____» ___________ 20__ год"
        )
        for cell in header_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        doc.add_paragraph("1. " + ticket[0])
        doc.add_paragraph("2. " + ticket[1])
        doc.add_paragraph("3.* " + "______________________________________________________________")
        doc.add_paragraph()
        doc.add_paragraph("Преподаватель ___________________ " + teacher)
        doc.add_page_break()

    doc.save(output_path)

def main_gui():
    root = Tk()
    root.title("Генератор экзаменационных билетов")
    root.geometry("600x500")
    root.configure(bg="#f0f8ff")  # светлый фон (AliceBlue)

    style = ttk.Style(root)
    style.theme_use('clam')
    style.configure("TButton", font=("Helvetica", 12), padding=6)
    style.configure("TLabel", font=("Helvetica", 12), background="#f0f8ff")
    style.configure("TCombobox", font=("Helvetica", 12))

    header_frame = Frame(root, bg="#4682b4", height=80)
    header_frame.pack(fill="x")
    header_label = Label(header_frame, text="Генератор экзаменационных билетов",
                         font=("Helvetica", 20, "bold"), bg="#4682b4", fg="white")
    header_label.pack(expand=True)

    main_frame = Frame(root, bg="#f0f8ff", pady=20)
    main_frame.pack(expand=True, fill="both")


    discipline_var = StringVar(value="Операционные системы")
    specialty_var = StringVar(value="02.03.02 ФИиИТ")
    group_var = StringVar(value="БФИ2102")

    Label(main_frame, text="Дисциплина:", bg="#f0f8ff").pack(pady=(10, 0))
    discipline_entry = Entry(main_frame, textvariable=discipline_var, font=("Helvetica", 12))
    discipline_entry.pack(pady=(0, 10))

    Label(main_frame, text="Специальность:", bg="#f0f8ff").pack(pady=(10, 0))
    specialty_entry = Entry(main_frame, textvariable=specialty_var, font=("Helvetica", 12))
    specialty_entry.pack(pady=(0, 10))

    Label(main_frame, text="Группа:", bg="#f0f8ff").pack(pady=(10, 0))
    group_entry = Entry(main_frame, textvariable=group_var, font=("Helvetica", 12))
    group_entry.pack(pady=(0, 10))


    format_var = StringVar(value="DOCX")
    format_label = Label(main_frame, text="Выберите формат генерации:", bg="#f0f8ff")
    format_label.pack(pady=(10, 5))
    format_combo = ttk.Combobox(main_frame, textvariable=format_var, values=["DOCX", "PDF"], state="readonly")
    format_combo.pack(pady=5)
    format_combo.current(0)

    def run_generation():
        try:
            use_api = messagebox.askyesno("Генерация вопросов",
                                          "Желаете сгенерировать вопросы по теме с использованием ИИ?\nЕсли нет, будут использованы вопросы из файла.")
            if use_api:
                topic = simpledialog.askstring("Тема", "Введите тему для генерации вопросов:")
                if not topic:
                    return
                num_tickets = simpledialog.askinteger("Количество билетов", "Сколько билетов нужно сгенерировать?")
                if not num_tickets:
                    return
                required = num_tickets * 2 * 2
                questions = generate_questions_by_topic(topic, required)
            else:
                input_path = filedialog.askopenfilename(title="Выберите .docx с вопросами",
                                                        filetypes=[("DOCX files", "*.docx")])
                if not input_path:
                    return
                questions = read_questions_from_docx(input_path)
                num_tickets = simpledialog.askinteger("Количество билетов", "Сколько билетов нужно сгенерировать?")
                if not num_tickets:
                    return

            tickets = generate_tickets_genetic(questions, num_tickets)
            now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
            output_dir = f"tickets_output_{now}"
            os.makedirs(output_dir, exist_ok=True)


            discipline = discipline_var.get()
            specialty = specialty_var.get()
            group = group_var.get()

            format_choice = format_var.get()
            if format_choice == "DOCX":
                output_file = os.path.join(output_dir, "bilety.docx")
                create_formatted_exam_docx(tickets, output_file, discipline=discipline, specialty=specialty, group=group)
                messagebox.showinfo("Успех", f"Файл создан:\n{output_file}")
            else:
                tex_file = os.path.join(output_dir, "bilety.tex")
                generate_latex(tickets, tex_file, discipline, specialty, group)
                compiled = compile_pdf(tex_file, output_dir)
                if compiled:
                    messagebox.showinfo("Успех", f"PDF сгенерирован в папке:\n{output_dir}")
                else:
                    messagebox.showwarning("Ошибка", "Ошибка при компиляции PDF. Убедитесь, что установлен pdflatex.")
        except Exception as e:
            messagebox.showerror("Ошибка", str(e))

    generate_button = ttk.Button(main_frame, text="Выбрать режим и сгенерировать", command=run_generation)
    generate_button.pack(pady=(20, 10), ipadx=10, ipady=5)

    root.mainloop()

if __name__ == "__main__":
    main_gui()
